import os
import re
import glob
import time
import base64
import zlib
import subprocess
import urllib.request
import urllib.error
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import markdown
from bs4 import BeautifulSoup, NavigableString
from PIL import Image

try:
    from lxml import etree as _lxml_etree
    LXML_AVAILABLE = True
except ImportError:
    LXML_AVAILABLE = False

# Cohesive corporate color palette (Classic Navy Theme)
PRIMARY_COLOR = RGBColor(0x1F, 0x4E, 0x79)    # Navy Blue (Headings)
SECONDARY_COLOR = RGBColor(0x5B, 0x9B, 0xD5)  # Light Blue
TEXT_COLOR = RGBColor(0x33, 0x33, 0x33)       # Charcoal/Off-black for body text
MUTED_COLOR = RGBColor(0x7F, 0x7F, 0x7F)      # Gray for captions
CODE_BG_HEX = "F2F2F2"                        # Light gray background for code blocks
TABLE_HEADER_BG_HEX = "E6EDF5"                # Light navy background for table headers
TABLE_BORDER_HEX = "D3D3D3"                   # Light gray borders

FONT_NAME = "Calibri"

# Unique sentinel — no markdown-special chars (**__~~`[]) so it survives the
# markdown renderer unchanged and we can detect it in the HTML output.
_MATH_SENTINEL = "zMATHBLK{:04d}z"
_MATH_SENTINEL_RE = re.compile(r'^zMATHBLK\d{4}z$')

def extract_math_blocks(text):
    """Extracts $$...$$ display math blocks before markdown processing.

    Returns the modified text (with placeholders) and a dict mapping
    each placeholder key → original LaTeX formula string."""
    blocks = {}
    parts = []
    pattern = re.compile(r'\$\$(.*?)\$\$', re.DOTALL)
    last_end = 0
    for i, m in enumerate(pattern.finditer(text)):
        key = _MATH_SENTINEL.format(i)
        blocks[key] = m.group(1).strip()
        parts.append(text[last_end:m.start()])
        # Surround with blank lines so markdown wraps it in its own <p>
        parts.append(f"\n\n{key}\n\n")
        last_end = m.end()
    parts.append(text[last_end:])
    return ''.join(parts), blocks

# ── LaTeX → native Word equation (OMML) ─────────────────────────────────────

_OMML_NS  = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_OMML_NSMAP = {'m': _OMML_NS}

# Command tables
_NARY_MAP = {'\\sum': '∑', '\\prod': '∏', '\\int': '∫',
             '\\bigcup': '⋃', '\\bigcap': '⋂'}
_SYM_MAP  = {
    '\\times': ' × ', '\\cdot': ' · ', '\\pm': ' ± ', '\\mp': ' ∓ ',
    '\\div': ' ÷ ', '\\infty': '∞',
    '\\leq': ' ≤ ', '\\geq': ' ≥ ', '\\neq': ' ≠ ', '\\approx': ' ≈ ',
    '\\le': ' ≤ ', '\\ge': ' ≥ ', '\\ne': ' ≠ ', '\\equiv': ' ≡ ',
    '\\alpha': 'α', '\\beta': 'β', '\\gamma': 'γ', '\\delta': 'δ',
    '\\Delta': 'Δ', '\\epsilon': 'ε', '\\varepsilon': 'ε',
    '\\theta': 'θ', '\\lambda': 'λ', '\\mu': 'μ', '\\nu': 'ν',
    '\\xi': 'ξ', '\\pi': 'π', '\\rho': 'ρ', '\\sigma': 'σ',
    '\\tau': 'τ', '\\phi': 'φ', '\\varphi': 'φ', '\\psi': 'ψ',
    '\\omega': 'ω', '\\Omega': 'Ω', '\\Sigma': 'Σ', '\\Pi': 'Π',
    '\\Lambda': 'Λ', '\\Theta': 'Θ',
    '\\%': '%', '\\$': '$', '\\#': '#', '\\&': '&',
    '\\,': ' ', '\\;': ' ', '\\!': '',
    '\\quad': '  ', '\\qquad': '    ',
    '\\cdots': '⋯', '\\ldots': '…', '\\dots': '…',
    '\\to': '→', '\\rightarrow': '→', '\\leftarrow': '←',
    '\\Rightarrow': '⇒', '\\Leftarrow': '⇐', '\\Leftrightarrow': '⟺',
    '\\in': ' ∈ ', '\\notin': ' ∉ ', '\\subset': ' ⊂ ', '\\supset': ' ⊃ ',
    '\\cup': ' ∪ ', '\\cap': ' ∩ ', '\\forall': '∀', '\\exists': '∃',
    '\\partial': '∂', '\\nabla': '∇',
}
_TEXT_CMDS = frozenset(['\\text', '\\mathrm', '\\mbox', '\\mathit',
                        '\\textrm', '\\textbf', '\\textnormal',
                        '\\mathnormal', '\\textit'])
_DELIM_MAP = {'\\{': '{', '\\}': '}', '\\|': '‖',
              '\\lfloor': '⌊', '\\rfloor': '⌋',
              '\\lceil': '⌈', '\\rceil': '⌉',
              '\\langle': '⟨', '\\rangle': '⟩'}

# ── element builders ──────────────────────────────────────────────────────────

def _mel(tag):
    return _lxml_etree.Element(f"{{{_OMML_NS}}}{tag}")

def _mse(parent, tag):
    return _lxml_etree.SubElement(parent, f"{{{_OMML_NS}}}{tag}")

def _mset(el, attr, val):
    el.set(f"{{{_OMML_NS}}}{attr}", val)

def _m_run(text, upright=False):
    r = _mel('r')
    if upright:
        rPr = _mse(r, 'rPr')
        _mse(rPr, 'nor')
    t = _mse(r, 't')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r

def _m_frac(num_els, den_els):
    f = _mel('f')
    n = _mse(f, 'num')
    for e in num_els: n.append(e)
    d = _mse(f, 'den')
    for e in den_els: d.append(e)
    return f

def _m_ssub(base, sub):
    s = _mel('sSub')
    be = _mse(s, 'e')
    for e in base: be.append(e)
    se = _mse(s, 'sub')
    for e in sub:  se.append(e)
    return s

def _m_ssup(base, sup):
    s = _mel('sSup')
    be = _mse(s, 'e')
    for e in base: be.append(e)
    se = _mse(s, 'sup')
    for e in sup:  se.append(e)
    return s

def _m_ssubsup(base, sub, sup):
    s = _mel('sSubSup')
    be = _mse(s, 'e')
    for e in base: be.append(e)
    se = _mse(s, 'sub')
    for e in sub:  se.append(e)
    ue = _mse(s, 'sup')
    for e in sup:  ue.append(e)
    return s

def _m_nary(chr_val, sub_els, sup_els, body_els):
    nary = _mel('nary')
    naryPr = _mse(nary, 'naryPr')
    chr_e = _mse(naryPr, 'chr');  _mset(chr_e, 'val', chr_val)
    ll    = _mse(naryPr, 'limLoc'); _mset(ll, 'val', 'undOvr')
    if not sub_els:
        sh = _mse(naryPr, 'subHide'); _mset(sh, 'val', '1')
    if not sup_els:
        sh = _mse(naryPr, 'supHide'); _mset(sh, 'val', '1')
    sub = _mse(nary, 'sub')
    for e in sub_els: sub.append(e)
    sup = _mse(nary, 'sup')
    for e in sup_els: sup.append(e)
    body = _mse(nary, 'e')
    for e in body_els: body.append(e)
    return nary

def _m_delim(beg, end, inner_els):
    d = _mel('d')
    dPr = _mse(d, 'dPr')
    bc  = _mse(dPr, 'begChr'); _mset(bc, 'val', beg)
    ec  = _mse(dPr, 'endChr'); _mset(ec, 'val', end)
    e = _mse(d, 'e')
    for x in inner_els: e.append(x)
    return d

# ── tokenizer ─────────────────────────────────────────────────────────────────

def _tokenize_latex(formula):
    tokens = []
    i, n = 0, len(formula)
    while i < n:
        c = formula[i]
        if c == '\\':
            j = i + 1
            if j < n:
                if formula[j].isalpha():
                    while j < n and formula[j].isalpha():
                        j += 1
                    tokens.append(('CMD', formula[i:j]))
                else:
                    tokens.append(('CMD', formula[i:j + 1]))
                    j += 1
            i = j
        elif c == '{':  tokens.append(('LB', None)); i += 1
        elif c == '}':  tokens.append(('RB', None)); i += 1
        elif c == '_':  tokens.append(('SUB', None)); i += 1
        elif c == '^':  tokens.append(('SUP', None)); i += 1
        elif c in (' ', '\t', '\n'): i += 1
        else:           tokens.append(('CH', c)); i += 1
    return tokens

# ── parser ────────────────────────────────────────────────────────────────────

class _LatexToOMML:
    def __init__(self, tokens):
        self.t   = tokens
        self.pos = 0

    def _peek(self):
        return self.t[self.pos] if self.pos < len(self.t) else None

    def _next(self):
        v = self.t[self.pos]; self.pos += 1; return v

    # ── public entry ──────────────────────────────────────────────────────────
    def parse(self, stop_types=None, stop_cmd=None):
        result = []
        while self.pos < len(self.t):
            tok = self._peek()
            if tok is None: break
            if stop_types and tok[0] in stop_types: break
            if stop_cmd  and tok[0] == 'CMD' and tok[1] == stop_cmd: break
            if tok[0] == 'RB': break

            elems = self._atom()
            if elems is None:
                self._next(); continue

            is_nary = (len(elems) == 1 and
                       elems[0].tag == f"{{{_OMML_NS}}}nary")
            sub_e = sup_e = None

            while self._peek() and self._peek()[0] in ('SUB', 'SUP'):
                kind   = self._next()[0]
                script = self._one()
                if is_nary:
                    nary   = elems[0]
                    naryPr = nary.find(f"{{{_OMML_NS}}}naryPr")
                    slot   = nary.find(f"{{{_OMML_NS}}}{'sub' if kind=='SUB' else 'sup'}")
                    for x in script: slot.append(x)
                    hide_tag = f"{{{_OMML_NS}}}{'subHide' if kind=='SUB' else 'supHide'}"
                    hide = naryPr.find(hide_tag)
                    if hide is not None: naryPr.remove(hide)
                else:
                    if kind == 'SUB': sub_e = script
                    else:             sup_e = script

            if is_nary:
                result.extend(elems)
            elif sub_e is not None and sup_e is not None:
                result.append(_m_ssubsup(elems, sub_e, sup_e))
            elif sub_e is not None:
                result.append(_m_ssub(elems, sub_e))
            elif sup_e is not None:
                result.append(_m_ssup(elems, sup_e))
            else:
                result.extend(elems)
        return result

    # ── helpers ───────────────────────────────────────────────────────────────
    def _one(self):
        """Single atom or {group}."""
        tok = self._peek()
        if not tok: return []
        if tok[0] == 'LB':
            self._next()
            elems = self.parse(stop_types={'RB'})
            if self._peek() and self._peek()[0] == 'RB': self._next()
            return elems
        return self._atom() or []

    def _atom(self):
        """One math atom → list of OMML elements (or None to skip)."""
        tok = self._peek()
        if not tok: return None

        if tok[0] == 'CH':
            self._next()
            ch = tok[1]
            if ch in ('=', '+', '-', '<', '>'): return [_m_run(f' {ch} ')]
            if ch in (',', ';', ':'):            return [_m_run(ch, upright=True)]
            return [_m_run(ch)]

        if tok[0] == 'CMD': return self._cmd()

        if tok[0] == 'LB':
            self._next()
            elems = self.parse(stop_types={'RB'})
            if self._peek() and self._peek()[0] == 'RB': self._next()
            return elems

        return None

    def _cmd(self):
        """Dispatch on command name → list of OMML elements."""
        cmd = self._next()[1]

        if cmd in _TEXT_CMDS:
            text = self._group_text()
            return [_m_run(text, upright=True)] if text else []

        if cmd == '\\frac':
            return [_m_frac(self._one(), self._one())]

        if cmd in _NARY_MAP:
            body = self._nary_body()
            return [_m_nary(_NARY_MAP[cmd], [], [], body)]

        if cmd == '\\left':
            beg   = self._delim_char()
            inner = self.parse(stop_cmd='\\right')
            if self._peek() and self._peek()[0]=='CMD' and self._peek()[1]=='\\right':
                self._next()
            end = self._delim_char()
            return [_m_delim(beg, end, inner)]

        if cmd == '\\right': return []

        if cmd == '\\sqrt':
            rad  = _mel('rad')
            _mse(rad, 'deg')
            body_el = _mse(rad, 'e')
            for x in self._one(): body_el.append(x)
            return [rad]

        if cmd in _SYM_MAP:
            sym = _SYM_MAP[cmd]
            return [_m_run(sym, upright=True)] if sym else []

        return [_m_run(cmd[1:], upright=True)]

    def _group_text(self):
        """Parse {content} → plain string (for \\text{}, \\mathrm{} etc.)."""
        if not self._peek() or self._peek()[0] != 'LB':
            if self._peek() and self._peek()[0] == 'CH':
                return self._next()[1]
            return ''
        self._next()  # {
        parts = []
        while self._peek() and self._peek()[0] != 'RB':
            t = self._next()
            if   t[0] == 'CH':  parts.append(t[1])
            elif t[0] == 'CMD':
                if t[1] in _TEXT_CMDS: parts.append(self._group_text())
                else:                  parts.append(t[1][1:])
        if self._peek() and self._peek()[0] == 'RB': self._next()  # }
        return ''.join(parts)

    def _delim_char(self):
        """Char immediately after \\left / \\right."""
        tok = self._peek()
        if not tok: return ''
        if tok[0] == 'CH':
            c = self._next()[1]
            return '' if c == '.' else c
        if tok[0] == 'CMD':
            return _DELIM_MAP.get(self._next()[1], '')
        return ''

    def _nary_body(self):
        """Body of \\sum / \\prod etc. – grab the next \\left…\\right or group."""
        tok = self._peek()
        if not tok: return []
        if tok[0] == 'CMD' and tok[1] == '\\left':
            return self._atom() or []
        if tok[0] == 'LB':
            return self._one()
        return self._atom() or []


def latex_to_omml_para(formula):
    """Convert a LaTeX formula string to an OMML <m:oMathPara> lxml element
    suitable for direct insertion into a python-docx w:p element."""
    tokens = _tokenize_latex(formula)
    elems  = _LatexToOMML(tokens).parse()
    para   = _lxml_etree.Element(f"{{{_OMML_NS}}}oMathPara", nsmap=_OMML_NSMAP)
    math   = _lxml_etree.SubElement(para, f"{{{_OMML_NS}}}oMath")
    for e in elems:
        math.append(e)
    return para

def set_cell_background(cell, fill_hex):
    """Sets the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_borders(cell, border_hex):
    """Sets light borders for table cells."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="{border_hex}"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="{border_hex}"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border_hex}"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="{border_hex}"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets margins (padding) inside table cells in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_left_indent(table, indent_value_dxa):
    """Sets left indentation for a table in twentieths of a point (dxa)."""
    tblPr = table._tbl.tblPr
    tblInd = parse_xml(f'<w:tblInd {nsdecls("w")} w:w="{indent_value_dxa}" w:type="dxa"/>')
    tblPr.append(tblInd)

def set_table_width(table, width_dxa):
    """Sets the explicit total width for a table in twentieths of a point (dxa)."""
    tblPr = table._tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>')
    tblPr.append(tblW)

def add_hyperlink(paragraph, url, text, color=None, underline=True):
    """Places a hyperlink within a paragraph using direct OpenXML manipulation."""
    part = paragraph.part
    # Register the link relationship
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink XML element
    hyperlink = parse_xml(
        f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        f'r:id="{r_id}"/>'
    )

    # Create a new run inside the hyperlink
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set formatting properties
    if color:
        color_hex = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
        rPr.append(parse_xml(f'<w:color {nsdecls("w")} w:val="{color_hex}"/>'))
        
    if underline:
        rPr.append(parse_xml(f'<w:u {nsdecls("w")} w:val="single"/>'))
        
    # Font properties
    rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'))
    rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="22"/>')) # 11pt
    
    new_run.append(rPr)
    
    # Add text node
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def format_run(run, font_name=FONT_NAME, size_pt=11, bold=False, italic=False, color=TEXT_COLOR):
    """Utility to easily format a text run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color

def parse_inline_elements(parent_node, paragraph_docx, is_code=False):
    """Traverses children of an HTML element to parse and apply inline formatting (bold, italic, code)."""
    for child in parent_node.contents:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip() or text == " ":
                run = paragraph_docx.add_run(text)
                if is_code:
                    format_run(run, italic=True)
                else:
                    format_run(run)
        elif child.name in ['strong', 'b']:
            run = paragraph_docx.add_run(child.get_text())
            format_run(run, bold=True)
        elif child.name in ['em', 'i']:
            run = paragraph_docx.add_run(child.get_text())
            format_run(run, italic=True)
        elif child.name in ['code']:
            run = paragraph_docx.add_run(child.get_text())
            format_run(run, italic=True)
        elif child.name == 'a':
            # Handle links as colored text
            run = paragraph_docx.add_run(child.get_text())
            format_run(run, color=SECONDARY_COLOR, italic=True)
        elif child.name == 'br':
            paragraph_docx.add_run('\n')

def add_header(doc, text, level):
    """Adds a stylish, color-branded heading with specific top/bottom margins."""
    # Headings mapping
    size_map = {1: 18, 2: 15, 3: 13, 4: 11.5}
    size = size_map.get(level, 11)
    
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(14 if level > 1 else 18)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True
    
    run = heading.add_run(text)
    format_run(run, size_pt=size, bold=True, color=PRIMARY_COLOR)

def build_docx_from_html(doc, soup_body, input_dir, output_dir=None, math_blocks=None):
    """Iterates through top-level HTML nodes and inserts them structured into the docx document."""

    for element in soup_body.contents:
        if isinstance(element, NavigableString):
            continue

        # Skip decomposed or detached elements (like consumed figure captions)
        if not element.parent:
            continue

        # 1. Headers (h1 - h6)
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            add_header(doc, element.get_text(), level)

        # 2. Paragraphs (p)
        elif element.name == 'p':
            # --- Display math block ($$...$$) ---
            p_text = element.get_text().strip()
            if math_blocks and _MATH_SENTINEL_RE.match(p_text):
                formula = math_blocks.get(p_text)
                if formula:
                    inserted = False
                    if LXML_AVAILABLE:
                        try:
                            omml_para = latex_to_omml_para(formula)
                            eq_p = doc.add_paragraph()
                            eq_p.paragraph_format.space_before = Pt(10)
                            eq_p.paragraph_format.space_after  = Pt(10)
                            eq_p._p.append(omml_para)
                            inserted = True
                        except Exception as _e:
                            print(f"   [WARNING] OMML insert failed: {_e}")
                    if not inserted:
                        # Plain-text fallback
                        fb_p = doc.add_paragraph()
                        fb_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        fb_p.paragraph_format.space_before = Pt(8)
                        fb_p.paragraph_format.space_after  = Pt(8)
                        run = fb_p.add_run(formula)
                        format_run(run, font_name="Consolas", size_pt=10, italic=True,
                                   color=RGBColor(0x3F, 0x3F, 0x3F))
                continue

            # Check if paragraph only contains a single image (allowing surrounding wrappers/whitespace)
            img_tags = element.find_all('img')
            img_tag = None
            if len(img_tags) == 1 and not element.get_text().strip():
                img_tag = img_tags[0]

            if img_tag:
                # Isolated block image
                # Check if the next sibling tag is a Figure caption
                custom_caption = None
                next_sibling = element.find_next_sibling()
                if next_sibling and next_sibling.name == 'p':
                    next_text = next_sibling.get_text().strip()
                    if next_text.lower().startswith("figure"):
                        custom_caption = next_text
                        next_sibling.decompose()  # Remove it so it is not processed as a paragraph
                
                process_image(doc, img_tag, input_dir, custom_caption=custom_caption, output_dir=output_dir)
            else:
                p_docx = doc.add_paragraph()
                p_docx.paragraph_format.left_indent = Inches(0.3)
                p_docx.paragraph_format.space_before = Pt(0)
                p_docx.paragraph_format.space_after = Pt(6)
                p_docx.paragraph_format.line_spacing = 1.15
                
                # Check for images inside paragraphs along with text
                for child in element.contents:
                    if child.name == 'img':
                        process_image(doc, child, input_dir, output_dir=output_dir)
                    else:
                        parse_inline_elements(element, p_docx)
                        break # parse_inline_elements will handle the whole text node

        # 3. Blockquotes (blockquote)
        elif element.name == 'blockquote':
            p_docx = doc.add_paragraph()
            p_docx.paragraph_format.left_indent = Inches(0.7)
            p_docx.paragraph_format.right_indent = Inches(0.4)
            p_docx.paragraph_format.space_before = Pt(6)
            p_docx.paragraph_format.space_after = Pt(6)
            
            # Apply italic and muted coloring
            for p_child in element.find_all('p'):
                parse_inline_elements(p_child, p_docx)
            # If plain content in blockquote
            if not element.find('p'):
                parse_inline_elements(element, p_docx)
            for run in p_docx.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # 4. Unordered & Ordered Lists (ul, ol)
        elif element.name in ['ul', 'ol']:
            style = 'List Bullet' if element.name == 'ul' else 'List Number'
            
            for idx, li in enumerate(element.find_all('li', recursive=False)):
                li_p = doc.add_paragraph(style=style)
                li_p.paragraph_format.space_before = Pt(0)
                li_p.paragraph_format.space_after = Pt(3)
                li_p.paragraph_format.line_spacing = 1.15
                
                # If there are sub-lists nested inside this li
                nested_list = li.find(['ul', 'ol'], recursive=False)
                
                # Parse the direct inline content of the list item
                # Extract text and child elements, excluding the nested list
                temp_span = BeautifulSoup("<span></span>", "html.parser").span
                for li_child in li.contents:
                    if li_child != nested_list:
                        temp_span.append(li_child.__copy__())
                
                parse_inline_elements(temp_span, li_p)
                
                # If a nested list is found, render it recursively
                if nested_list:
                    # Let's indent it manually
                    nested_style = 'List Bullet 2' if nested_list.name == 'ul' else 'List Number 2'
                    for sub_li in nested_list.find_all('li', recursive=False):
                        sub_p = doc.add_paragraph(style=nested_style)
                        sub_p.paragraph_format.space_before = Pt(0)
                        sub_p.paragraph_format.space_after = Pt(3)
                        sub_p.paragraph_format.line_spacing = 1.15
                        parse_inline_elements(sub_li, sub_p)

        # 5. Tables (table)
        elif element.name == 'table':
            rows = element.find_all('tr')
            if not rows:
                continue
                
            # Determine maximum column count
            col_count = 0
            for row in rows:
                cells = row.find_all(['td', 'th'])
                col_count = max(col_count, len(cells))
                
            if col_count == 0:
                continue
                
            # Create a styled Word table
            table_docx = doc.add_table(rows=0, cols=col_count)
            table_docx.autofit = False
            table_docx.allow_autofit = False
            
            # Set left indentation of 0.3" (432 dxa) and maximum width of 6.0" (8640 dxa)
            set_table_left_indent(table_docx, 432)
            set_table_width(table_docx, 8640)
            
            for r_idx, row in enumerate(rows):
                row_docx = table_docx.add_row()
                cells = row.find_all(['td', 'th'])
                
                for c_idx, cell in enumerate(cells):
                    if c_idx >= col_count:
                        break
                    
                    docx_cell = row_docx.cells[c_idx]
                    docx_cell.width = Inches(6.0 / col_count)
                    
                    cell_p = docx_cell.paragraphs[0]
                    cell_p.paragraph_format.space_before = Pt(2)
                    cell_p.paragraph_format.space_after = Pt(2)
                    
                    # Fill cell with inline formatted text
                    is_header = (cell.name == 'th' or r_idx == 0)
                    parse_inline_elements(cell, cell_p)
                    
                    # Style modifications
                    set_cell_borders(docx_cell, TABLE_BORDER_HEX)
                    set_cell_margins(docx_cell, top=120, bottom=120, left=180, right=180)
                    
                    if is_header:
                        # Make header text bold
                        for run in cell_p.runs:
                            run.bold = True
                            run.font.color.rgb = PRIMARY_COLOR
                        # Add a soft gray background
                        set_cell_background(docx_cell, TABLE_HEADER_BG_HEX)
                    else:
                        # Regular rows
                        if r_idx % 2 == 0:
                            # Light zebra striping
                            set_cell_background(docx_cell, "F9F9F9")
            
            # Add small spacing after tables
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(4)
            spacer.paragraph_format.space_after = Pt(4)

        # 6. Preformatted/Code blocks (pre)
        elif element.name == 'pre':
            code_tag = element.find('code')
            code_text = code_tag.get_text() if code_tag else element.get_text()
            
            p_docx = doc.add_paragraph()
            p_docx.paragraph_format.left_indent = Inches(0.6)
            p_docx.paragraph_format.space_before = Pt(6)
            p_docx.paragraph_format.space_after = Pt(6)
            p_docx.paragraph_format.line_spacing = 1.0
            
            # Format block code background and border
            pPr = p_docx._p.get_or_add_pPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG_HEX}"/>')
            pPr.append(shd)
            
            # Add text
            run = p_docx.add_run(code_text.rstrip())
            format_run(run, font_name="Consolas", size_pt=9, color=RGBColor(0x3F, 0x3F, 0x3F))

        # 7. Horizontal Rules (hr) - Skipped as per user request
        elif element.name == 'hr':
            pass

def process_image(doc, img_tag, input_dir, custom_caption=None, output_dir=None):
    """Resolves local images or attachments, scales images cleanly, and embeds them with captions or attachment cards."""
    src = img_tag.get('src', '')
    alt = img_tag.get('alt', 'Embedded Screenshot')
    
    # Resolve the file path relative to the input folder
    file_path = os.path.abspath(os.path.join(input_dir, src))
    filename = os.path.basename(src)
    
    if not os.path.exists(file_path):
        # Fallback to direct check in input_dir
        fallback_path = os.path.join(input_dir, filename)
        if os.path.exists(fallback_path):
            file_path = fallback_path
            
    # Determine extension
    ext = os.path.splitext(filename.lower())[1]
    is_image = ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp']
    
    # If the extension is .x-emf, it is a vector OLE spreadsheet placeholder
    if ext == '.x-emf':
        # Check if there is an associated excel spreadsheet
        xlsx_filename = filename.replace('.x-emf', '.xlsx')
        xlsx_path = os.path.join(input_dir, xlsx_filename)
        if os.path.exists(xlsx_path):
            filename = xlsx_filename
            file_path = xlsx_path
            ext = '.xlsx'
            is_image = False
            
    if not is_image:
        # Non-image file attachment! (e.g. .xlsx, .pdf, .zip, etc.)
        # Determine icon/prefix based on extension
        if ext in ['.xlsx', '.xls']:
            file_icon = "📊 [Excel Spreadsheet]"
        elif ext in ['.docx', '.doc']:
            file_icon = "📝 [Word Document]"
        elif ext in ['.pdf']:
            file_icon = "📕 [PDF Document]"
        elif ext in ['.zip', '.rar', '.7z']:
            file_icon = "📦 [Compressed Archive]"
        elif ext in ['.csv']:
            file_icon = "📋 [CSV Spreadsheet]"
        else:
            file_icon = "📎 [File Attachment]"
            
        # Copy the file to output directory so the relative hyperlink works
        if output_dir:
            import shutil
            os.makedirs(output_dir, exist_ok=True)
            dest_path = os.path.join(output_dir, filename)
            try:
                shutil.copy2(file_path, dest_path)
                print(f"   [ATTACH] Copied file attachment to output: {filename}")
            except Exception as copy_err:
                print(f"   [WARNING] Could not copy attachment {filename}: {str(copy_err)}")
                
        # Create a gorgeous, premium styled Card block in the document!
        p_docx = doc.add_paragraph()
        p_docx.paragraph_format.left_indent = Inches(0.3)
        p_docx.paragraph_format.space_before = Pt(8)
        p_docx.paragraph_format.space_after = Pt(8)
        p_docx.paragraph_format.keep_with_next = True
        
        # Soft background color (light blue for sheets, light gray for others)
        pPr = p_docx._p.get_or_add_pPr()
        bg_color = "E6EDF5" if ext in ['.xlsx', '.xls', '.csv'] else "F2F2F2"
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
        pPr.append(shd)
        
        # Border decorations
        borders = parse_xml(f'''
            <w:pBdr {nsdecls("w")}>
                <w:top w:val="single" w:sz="4" w:space="6" w:color="D3D3D3"/>
                <w:left w:val="single" w:sz="16" w:space="10" w:color="1F4E79"/>
                <w:bottom w:val="single" w:sz="4" w:space="6" w:color="D3D3D3"/>
                <w:right w:val="single" w:sz="4" w:space="6" w:color="D3D3D3"/>
            </w:pBdr>
        ''')
        pPr.append(borders)
        
        # Prefix icon
        prefix_run = p_docx.add_run(f"  {file_icon}   ")
        format_run(prefix_run, bold=True, color=PRIMARY_COLOR)
        
        # Clickable Hyperlink
        add_hyperlink(p_docx, filename, filename, color=SECONDARY_COLOR, underline=True)
        
        # If it is a spreadsheet, also insert a shape placeholder so OLE embedding post-processing can run!
        if ext in ['.xlsx', '.xls']:
            # Create a 1x1 pixel white png placeholder in memory
            placeholder_path = os.path.join(input_dir, "attachment_ole_placeholder.png")
            if not os.path.exists(placeholder_path):
                try:
                    img = Image.new('RGB', (1, 1), color='white')
                    img.save(placeholder_path)
                except:
                    pass
            
            if os.path.exists(placeholder_path):
                # Add it to the paragraph
                run = p_docx.add_run("   ")
                try:
                    picture = run.add_picture(placeholder_path, width=Inches(0.1), height=Inches(0.1))
                    # Set alt text / description to the Excel filename
                    docPr = picture._inline.docPr
                    docPr.set('name', filename)
                    docPr.set('descr', filename)
                    docPr.set('title', filename)
                except Exception as shape_err:
                    print(f"   [WARNING] Could not insert OLE shape placeholder: {str(shape_err)}")
                    
        # Create elegant caption block below the attachment block
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.paragraph_format.space_before = Pt(4)
        caption_p.paragraph_format.space_after = Pt(10)
        
        if custom_caption:
            match = re.match(r'^figure[:\s]*(.*)', custom_caption, re.IGNORECASE)
            if match:
                rest = match.group(1).strip()
                if rest:
                    normalized_caption = f"Figure: {rest}"
                else:
                    normalized_caption = "Figure"
            else:
                normalized_caption = custom_caption
            caption_run = caption_p.add_run(normalized_caption)
        else:
            caption_run = caption_p.add_run(f"Figure: Attachment - {filename}")
            
        format_run(caption_run, size_pt=9, italic=True, color=MUTED_COLOR)
        return

    if os.path.exists(file_path):
        try:
            # 1. Load image to get width and height dimensions
            with Image.open(file_path) as img:
                width, height = img.size
                aspect_ratio = height / width
            
            # Max printable page width in Word (8.5 inches minus 2 inches margins = 6.5 inches)
            MAX_PRINT_WIDTH_INCHES = 6.0
            
            # Base sizing algorithm
            # Standard screen screenshot at ~96 DPI: width / 96
            inferred_width_inches = width / 110.0 # 110 DPI scale fits nicely
            
            if inferred_width_inches > MAX_PRINT_WIDTH_INCHES:
                target_width = MAX_PRINT_WIDTH_INCHES
            elif inferred_width_inches < 2.0:
                # Keep small icons/decorations small, but maybe bump up slightly
                target_width = max(inferred_width_inches, 1.5)
            else:
                target_width = inferred_width_inches
                
            # Create centered image block
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_p.paragraph_format.space_before = Pt(8)
            img_p.paragraph_format.space_after = Pt(4)
            img_p.paragraph_format.keep_with_next = True
            
            run = img_p.add_run()
            run.add_picture(file_path, width=Inches(target_width))
            
            # Create elegant caption block below the image
            caption_p = doc.add_paragraph()
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_p.paragraph_format.space_before = Pt(0)
            caption_p.paragraph_format.space_after = Pt(10)
            
            if custom_caption:
                # Normalize custom captions to ensure standard "Figure: [Text]" format
                match = re.match(r'^figure[:\s]*(.*)', custom_caption, re.IGNORECASE)
                if match:
                    rest = match.group(1).strip()
                    if rest:
                        normalized_caption = f"Figure: {rest}"
                    else:
                        normalized_caption = "Figure"
                else:
                    normalized_caption = custom_caption
                caption_run = caption_p.add_run(normalized_caption)
            else:
                caption_run = caption_p.add_run(f"Figure: {alt}")
            format_run(caption_run, size_pt=9, italic=True, color=MUTED_COLOR)
            
        except Exception as e:
            # Error reading image file
            error_p = doc.add_paragraph()
            run = error_p.add_run(f"[Error loading image {src}: {str(e)}]")
            format_run(run, bold=True, color=RGBColor(0xC0, 0x00, 0x00))
    else:
        # Image path not found
        filename = os.path.basename(src).lower()
        if "_add.png" in filename or "_detail.png" in filename:
            # Silently skip optional missing screenshots (Add/Detail forms)
            return
            
        missing_p = doc.add_paragraph()
        run = missing_p.add_run(f"[Missing Image Placeholder: \"{src}\" not found in input folder]")
        format_run(run, bold=True, color=RGBColor(0xC0, 0x00, 0x00))

def post_process_ole_embeddings(docx_path, input_dir):
    """
    Post-processes the generated .docx file on Windows using Word COM automation to
    replace any .x-emf image placeholders with actual embedded Excel (.xlsx) files if present.
    Stabilized with transient COM failure retries, Word auto-restart on RPC crash, and timing sleeps.
    """
    try:
        import win32com.client as win32
    except ImportError:
        print("\n[INFO] 'pywin32' is not installed. Skipping automatic Excel OLE embedding post-processing.")
        print("To enable automatic OLE embedding, run: pip install pywin32")
        return

    print("\n[INFO] Running post-processing OLE embedding using Word COM Automation...")
    word = None
    doc = None
    
    def start_word_and_open_doc(path):
        nonlocal word, doc
        # Clean up any existing instances first
        if doc:
            try: doc.Close(False)
            except: pass
        if word:
            try: word.Quit()
            except: pass
        # Start new Word COM instance with robust fallback to Dispatch if EnsureDispatch fails
        try:
            word = win32.gencache.EnsureDispatch('Word.Application')
        except Exception:
            try:
                word = win32.Dispatch('Word.Application')
            except Exception as dispatch_err:
                raise RuntimeError(f"Both EnsureDispatch and Dispatch failed to start Word. Details: {str(dispatch_err)}")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(os.path.abspath(path))

    try:
        start_word_and_open_doc(docx_path)
        
        shapes_to_replace = []
        
        for idx in range(1, doc.InlineShapes.Count + 1):
            shape = doc.InlineShapes(idx)
            alt = (shape.AlternativeText or "").strip()
            title = (shape.Title or "").strip()
            
            filename = None
            for term in [alt, title]:
                if "extracted_screenshot_" in term or term.lower().endswith(".xlsx"):
                    filename = term
                    break
                    
            if not filename:
                continue
                
            if filename.lower().endswith(".xlsx"):
                xlsx_path = os.path.join(input_dir, os.path.basename(filename))
            else:
                # Extract the base name (e.g. extracted_screenshot_20)
                base_name = os.path.splitext(os.path.basename(filename))[0]
                xlsx_filename = f"{base_name}.xlsx"
                xlsx_path = os.path.join(input_dir, xlsx_filename)
            
            if os.path.exists(xlsx_path):
                shapes_to_replace.append((idx, xlsx_path))
                
        # Replace shapes in reverse order to maintain index stability
        shapes_to_replace.reverse()
        for idx, xlsx_path in shapes_to_replace:
            # We retry each shape up to 3 times in case of transient RPC crashes!
            for attempt in range(1, 4):
                try:
                    shape = doc.InlineShapes(idx)
                    r = shape.Range
                    # Embed OLE object at the placeholder range first
                    r.InlineShapes.AddOLEObject(
                        FileName=os.path.abspath(xlsx_path),
                        DisplayAsIcon=False,
                        LinkToFile=False
                    )
                    # Now delete the original placeholder shape
                    shape.Delete()
                    print(f"   [EMBED] Embedded Excel worksheet: {os.path.basename(xlsx_path)}")
                    time.sleep(1.2)  # Pause to allow Word and background Excel processes to settle!
                    break # Success! Go to next shape
                except Exception as shape_err:
                    print(f"   [WARNING] Attempt {attempt} failed for OLE object {os.path.basename(xlsx_path)}: {str(shape_err)}")
                    if attempt < 3:
                        # RPC server might have crashed. Let's wait, restart Word, and re-open the document!
                        print("   [INFO] Re-initializing Word COM instance and retrying...")
                        time.sleep(2.0)
                        try:
                            start_word_and_open_doc(docx_path)
                        except Exception as init_err:
                            print(f"   [ERROR] Failed to re-initialize Word COM: {str(init_err)}")
                    else:
                        print(f"   [ERROR] Permanent failure embedding Excel OLE object {os.path.basename(xlsx_path)} after 3 attempts.")
            
        doc.Save()
        doc.Close()
        print("[OK] OLE post-processing complete!")
        
    except Exception as e:
        print(f"   [ERROR] Error during OLE post-processing: {str(e)}")
    finally:
        if word:
            try:
                word.Quit()
            except:
                pass

def render_mermaid_to_png(diagram_code, output_path):
    """Render a Mermaid diagram string to a PNG file.

    Tries mmdc (mermaid-cli) first; falls back to the kroki.io HTTP API.
    Returns True on success, False on failure.
    """
    import tempfile

    # --- attempt 1: mmdc CLI (requires Node.js + @mermaid-js/mermaid-cli) ---
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as tf:
            tf.write(diagram_code)
            mmd_path = tf.name

        result = subprocess.run(
            ['mmdc', '-i', mmd_path, '-o', output_path, '-b', 'white', '--width', '1200'],
            capture_output=True, text=True, timeout=30
        )
        try:
            os.unlink(mmd_path)
        except OSError:
            pass

        if result.returncode == 0 and os.path.exists(output_path):
            return True
        print(f"   [MERMAID] mmdc exited {result.returncode}: {result.stderr.strip()}")
    except FileNotFoundError:
        pass  # mmdc not installed
    except subprocess.TimeoutExpired:
        print("   [MERMAID] mmdc timed out")

    # --- attempt 2: kroki.io public API (no local install required) ---
    try:
        encoded = base64.urlsafe_b64encode(
            zlib.compress(diagram_code.encode('utf-8'), 9)
        ).decode('ascii')
        url = f"https://kroki.io/mermaid/png/{encoded}"
        req = urllib.request.Request(url, headers={'User-Agent': 'md-to-docx/1.0'})
        with urllib.request.urlopen(req, timeout=30) as resp:
            png_data = resp.read()
        with open(output_path, 'wb') as f:
            f.write(png_data)
        return True
    except Exception as e:
        print(f"   [MERMAID] kroki.io render failed: {e}")

    return False


def extract_mermaid_blocks(text, temp_dir):
    """Replace ```mermaid fences with rendered PNG image references.

    Renders each diagram to a temp PNG in temp_dir, replaces the fenced block
    with a Markdown image tag followed by a Figure caption paragraph.
    Returns (modified_text, list_of_temp_png_paths).
    """
    pattern = re.compile(r'```mermaid\s*\n(.*?)```', re.DOTALL | re.IGNORECASE)
    temp_files = []
    counter = [0]

    def replacer(m):
        diagram_code = m.group(1).strip()
        idx = counter[0]
        counter[0] += 1

        png_filename = f"_mermaid_diagram_{idx}.png"
        png_path = os.path.join(temp_dir, png_filename)

        print(f"   [MERMAID] Rendering diagram {idx + 1}...")
        if render_mermaid_to_png(diagram_code, png_path):
            temp_files.append(png_path)
            return (
                f"\n\n![Mermaid Diagram {idx + 1}]({png_filename})\n\n"
                f"Figure: Diagram {idx + 1}\n\n"
            )
        print(f"   [WARNING] Could not render Mermaid diagram {idx + 1}, inserting placeholder text.")
        return f"\n\n*[Mermaid diagram {idx + 1} could not be rendered]*\n\n"

    modified = pattern.sub(replacer, text)
    return modified, temp_files


def compile_markdown_to_docx(input_folder, output_filepath):
    """Assembles all markdown files alphabetically, processes image references, and builds a gorgeous .docx."""
    
    print(f"Reading Markdown files from: {input_folder}")
    
    # Grab all .md files in the input folder
    md_files = glob.glob(os.path.join(input_folder, "*.md"))
    # Sort alphabetically to guarantee consistent order
    md_files.sort()
    
    if not md_files:
        print(f"Warning: No .md files found in {input_folder}. Generating empty template document.")
        # Create standard empty doc
        doc = Document()
        doc.add_heading("Empty Document Template", level=1)
        doc.add_paragraph("Place your .md files and screenshots in the input folder and run the script.")
        doc.save(output_filepath)
        return
        
    # Initialize premium document structure
    doc = Document()
    
    # Configure 1-inch standard margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    print(f"Discovered {len(md_files)} markdown files. Compiling...")

    all_mermaid_temps = []

    # Process each markdown file sequentially
    for idx, filepath in enumerate(md_files):
        filename = os.path.basename(filepath)
        print(f" Processing: {filename}")

        with open(filepath, 'r', encoding='utf-8') as f:
            md_text = f.read()

        # Character normalization replacements
        md_text = md_text.replace('\u2018', "'").replace('\u2019', "'")
        md_text = md_text.replace('\u201c', '"').replace('\u201d', '"')
        md_text = md_text.replace('\u2014', '-')
        md_text = md_text.replace('\u00d7', 'x')
        md_text = md_text.replace('\u2705', 'Yes')
        md_text = md_text.replace('\u274c', 'No')

        # Render ```mermaid blocks to PNG images BEFORE markdown processing.
        md_text, mermaid_temps = extract_mermaid_blocks(md_text, input_folder)
        all_mermaid_temps.extend(mermaid_temps)

        # Extract $$...$$ display math blocks BEFORE markdown processing so that
        # underscores and asterisks inside formulas are not interpreted as markup.
        md_text, math_blocks = extract_math_blocks(md_text)

        # Convert Markdown formatting to HTML with rich layout support
        html_text = markdown.markdown(md_text, extensions=['extra', 'tables', 'fenced_code', 'nl2br'])
        soup = BeautifulSoup(html_text, 'html.parser')

        # If this is the second or subsequent file, optionally add a page break
        if idx > 0:
            doc.add_page_break()

        build_docx_from_html(doc, soup, input_folder,
                             output_dir=os.path.dirname(os.path.abspath(output_filepath)),
                             math_blocks=math_blocks)
        
    # Save the final file
    print(f"Saving compiled Word document to: {output_filepath}")
    saved_path = output_filepath
    try:
        doc.save(output_filepath)
        print("Compilation Complete!")
    except PermissionError:
        from datetime import datetime
        timestamp = datetime.now().strftime("%H%M%S")
        fallback_path = output_filepath.replace(".docx", f"_{timestamp}.docx")
        print(f"\n[ERROR] Permission Denied! The file is locked (possibly open in Microsoft Word).")
        print(f"Saving a backup copy as: {fallback_path}")
        doc.save(fallback_path)
        print("Backup Copy Compilation Complete!")
        saved_path = fallback_path

    # Run OLE post-processing step
    #post_process_ole_embeddings(saved_path, input_folder)

    # Clean up temporary Mermaid PNG files (already embedded inside the docx)
    for tmp_png in all_mermaid_temps:
        try:
            os.unlink(tmp_png)
        except OSError:
            pass

if __name__ == "__main__":
    # Standard directories
    base_dir = os.path.dirname(os.path.abspath(__file__))
    input_dir = os.path.join(base_dir, "input")
    output_dir = os.path.join(base_dir, "output")
    
    # Ensure folders exist
    os.makedirs(input_dir, exist_ok=True)
    os.makedirs(output_dir, exist_ok=True)
    
    output_docx = os.path.join(output_dir, "Functional_Specification_Document.docx")
    
    compile_markdown_to_docx(input_dir, output_docx)
